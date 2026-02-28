"""SAT XML Report Generator using Anthropic Claude API.

This module takes raw station data (tank readings, receptions, sales)
and uses Claude Opus to generate SAT-compliant XML for controles volumetricos.
"""
import os
import io
import json
import base64
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, date

try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")

# SAT Anexo 30 XML template for a gasolinera (expendio al publico)
SAT_XML_TEMPLATE = '''<?xml version="1.0" encoding="UTF-8"?>
<Covol:ControlVolumetrico
  xmlns:Covol="http://www.sat.gob.mx/ControlesVolumetricos"
  xmlns:Expendio="http://www.sat.gob.mx/ControlesVolumetricos/Expendio"
  xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance"
  xsi:schemaLocation="http://www.sat.gob.mx/ControlesVolumetricos https://repositorio.cloudb.sat.gob.mx/Covol/xml/Diario.xsd
                       http://www.sat.gob.mx/ControlesVolumetricos/Expendio https://repositorio.cloudb.sat.gob.mx/Covol/xml/Comp-Expendio.xsd"
  Covol:Version="1.0"
  Covol:RfcContribuyente="{rfc}"
  Covol:RfcRepresentanteLegal=""
  Covol:RfcProveedor="{rfc_proveedor}"
  Covol:Caracter="expendedor"
  Covol:ModalidadPermiso="{modalidad_permiso}"
  Covol:NumPermiso="{num_permiso}"
  Covol:NumContratoOAsignacion=""
  Covol:InstalacionAlmacenGasNatural=""
  Covol:ClaveInstalacion="{clave_instalacion}"
  Covol:DescripcionInstalacion="{descripcion_instalacion}"
  Covol:NumeroPozos="0"
  Covol:NumeroTanques="{num_tanques}"
  Covol:NumeroDuctosEntradaSalida="0"
  Covol:NumeroDuctosTransporteDistribucion="0"
  Covol:NumeroDispensarios="{num_dispensarios}"
  Covol:FechaYHoraCorte="{fecha_corte}">

  <Covol:Geolocalizacion
    Covol:GeolocalizacionLatitud="{latitud}"
    Covol:GeolocalizacionLongitud="{longitud}"/>

  {productos_xml}

  {bitacora_xml}

</Covol:ControlVolumetrico>'''


SYSTEM_PROMPT = """You are an expert Mexican fiscal compliance assistant specializing in SAT controles volumetricos (Anexo 30/21).

Your job is to take raw operational data from a gas station and generate a COMPLETE, VALID XML daily report following the SAT's Covol namespace schema.

CRITICAL RULES:
1. Use Covol: prefix for all control volumetrico elements
2. Use Expendio: prefix for all complemento elements (CFDI, Nacional, etc.)
3. All volumes in liters (UM03), prices in MXN
4. Balance MUST validate: VolumenExistencias = VolumenExistenciasAnterior + VolumenAcumOpsRecep - VolumenAcumOpsEntreg
5. Every reception must have a Complemento with Nacional/CFDI data
6. Every tank delivery must reference its Dispensario
7. Every dispensario delivery must have PrecioVentaTotEnt and CFDI
8. Product codes: PR09=Magna(87oct), PR07=Premium(91oct), PR03=Diesel
9. FechaYHoraCorte must be end of day: YYYY-MM-DDT23:59:59
10. Temperature range: 15-45Â°C, PresionAbsoluta: 0.0 for liquids

OUTPUT: Return ONLY the complete XML document. No markdown, no code fences, no explanation. Just the raw XML starting with <?xml and ending with </Covol:ControlVolumetrico>."""


EXTRACTION_SYSTEM_PROMPT = """Eres un experto en operaciones de estaciones de servicio (gasolineras) en Mexico. Tu trabajo es analizar documentos operativos y extraer datos estructurados para generar reportes SAT de controles volumetricos.

Analiza el documento proporcionado y extrae la siguiente informacion en formato JSON:

{
  "rfc": "RFC del contribuyente (si aparece)",
  "permiso": "Numero de permiso CRE/CNE (si aparece)",
  "clave_instalacion": "Clave de instalacion (si aparece)",
  "fecha": "Fecha del reporte en formato YYYY-MM-DD",
  "tanques": [
    {
      "nombre": "Nombre/ID del tanque",
      "producto": "magna|premium|diesel",
      "capacidad_litros": 40000,
      "inventario_inicial": 15000,
      "inventario_final": 16500,
      "temperatura": 26,
      "uncertain": false
    }
  ],
  "recepciones": [
    {
      "tanque": "Nombre/ID del tanque destino",
      "producto": "magna|premium|diesel",
      "litros": 20000,
      "proveedor": "Nombre del proveedor",
      "rfc_proveedor": "RFC del proveedor",
      "num_factura": "Numero de factura/CFDI",
      "precio_litro": 21.50,
      "fecha_hora": "2026-02-27T10:15:00",
      "uncertain": false
    }
  ],
  "entregas": [
    {
      "tanque": "Nombre/ID del tanque origen",
      "producto": "magna|premium|diesel",
      "litros": 18500,
      "dispensario": "ID del dispensario",
      "uncertain": false
    }
  ],
  "notas": [
    "Lista de observaciones, datos faltantes, o valores inciertos"
  ],
  "confidence": 85
}

REGLAS DE EXTRACCION:
1. Codigos de producto: Magna/Regular/87oct = "magna", Premium/91oct = "premium", Diesel = "diesel"
2. Si un valor no esta claro, marcalo con "uncertain": true y agrega una nota explicando
3. Si falta informacion critica (inventarios, recepciones), agrega una nota con "FALTANTE: ..."
4. El campo "confidence" es tu puntuacion de confianza (1-100):
   - Base: 100
   - Resta 5 por cada campo critico faltante (RFC, permiso, inventarios)
   - Resta 3 por cada valor incierto/estimado
   - Resta 2 por cada problema de calidad (documento borroso, datos ambiguos)
5. Volumenes siempre en litros, precios en MXN
6. Si el documento es una imagen borrosa o de baja calidad, refleja esto en el confidence score

OUTPUT: Responde SOLO con el JSON valido. Sin markdown, sin code fences, sin explicacion."""


def extract_text_from_file(file_bytes, filename):
    """Extract text content from uploaded file based on its type.

    Args:
        file_bytes: raw bytes of the uploaded file
        filename: original filename (used to determine type)

    Returns:
        dict with 'text' (extracted text) or 'image_base64' (for images),
        plus 'type' indicating the file type
    """
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext == 'pdf':
        if not HAS_PDFPLUMBER:
            return {"error": "pdfplumber no instalado. Contacte al administrador."}
        try:
            pdf = pdfplumber.open(io.BytesIO(file_bytes))
            pages_text = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables()
                table_text = ""
                for table in tables:
                    for row in table:
                        table_text += " | ".join(str(cell or "") for cell in row) + "\n"
                pages_text.append(f"--- Pagina {i+1} ---\n{text}\n{table_text}")
            pdf.close()
            return {"text": "\n".join(pages_text), "type": "pdf"}
        except Exception as e:
            return {"error": f"Error leyendo PDF: {str(e)}"}

    elif ext in ('xlsx', 'xls'):
        if not HAS_OPENPYXL:
            return {"error": "openpyxl no instalado."}
        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
            sheets_text = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                lines = [f"--- Hoja: {sheet_name} ---"]
                for row in ws.iter_rows(values_only=True):
                    line = " | ".join(str(cell if cell is not None else "") for cell in row)
                    if line.strip(" |"):
                        lines.append(line)
                sheets_text.append("\n".join(lines))
            wb.close()
            return {"text": "\n".join(sheets_text), "type": "xlsx"}
        except Exception as e:
            return {"error": f"Error leyendo Excel: {str(e)}"}

    elif ext == 'docx':
        if not HAS_DOCX:
            return {"error": "python-docx no instalado."}
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    paragraphs.append(" | ".join(cells))
            return {"text": "\n".join(paragraphs), "type": "docx"}
        except Exception as e:
            return {"error": f"Error leyendo DOCX: {str(e)}"}

    elif ext in ('jpg', 'jpeg', 'png', 'webp'):
        # For images, encode as base64 for Claude vision API
        mime_types = {'jpg': 'image/jpeg', 'jpeg': 'image/jpeg', 'png': 'image/png', 'webp': 'image/webp'}
        mime = mime_types.get(ext, 'image/jpeg')
        b64 = base64.standard_b64encode(file_bytes).decode('utf-8')
        return {"image_base64": b64, "mime_type": mime, "type": "image"}

    else:
        return {"error": f"Tipo de archivo no soportado: .{ext}"}


def extract_data_from_file(file_bytes, filename):
    """Use Claude to extract structured data from an uploaded document.

    Args:
        file_bytes: raw bytes of the uploaded file
        filename: original filename

    Returns:
        dict with extracted_data, confidence, notes, or error
    """
    if not HAS_ANTHROPIC:
        return {"error": "Anthropic SDK not installed."}
    if not ANTHROPIC_API_KEY:
        return {"error": "ANTHROPIC_API_KEY not configured."}

    # Extract content from file
    file_content = extract_text_from_file(file_bytes, filename)

    if "error" in file_content:
        return file_content

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        # Build message based on file type
        if file_content["type"] == "image":
            messages = [{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": file_content["mime_type"],
                            "data": file_content["image_base64"],
                        }
                    },
                    {
                        "type": "text",
                        "text": f"Analiza esta imagen de un documento operativo de gasolinera y extrae los datos estructurados. Archivo: {filename}"
                    }
                ]
            }]
        else:
            messages = [{
                "role": "user",
                "content": f"Analiza el siguiente documento operativo de gasolinera y extrae los datos estructurados.\n\nArchivo: {filename}\n\nCONTENIDO DEL DOCUMENTO:\n{file_content['text']}"
            }]

        message = client.messages.create(
            model="claude-opus-4-5-20251101",
            max_tokens=8000,
            system=EXTRACTION_SYSTEM_PROMPT,
            messages=messages,
        )

        response_text = message.content[0].text.strip()

        # Clean up markdown fences if present
        if response_text.startswith("```"):
            lines = response_text.split("\n")
            response_text = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

        # Parse JSON response
        extracted = json.loads(response_text)

        return {
            "success": True,
            "extracted_data": extracted,
            "confidence": extracted.get("confidence", 50),
            "notes": extracted.get("notas", []),
            "tokens_used": {
                "input": message.usage.input_tokens,
                "output": message.usage.output_tokens,
            },
        }

    except json.JSONDecodeError as e:
        return {"error": f"Error parseando respuesta de IA: {str(e)}", "raw_response": response_text}
    except anthropic.APIError as e:
        return {"error": f"Anthropic API error: {str(e)}"}
    except Exception as e:
        return {"error": f"Error inesperado: {str(e)}"}


def generate_sat_xml_with_ai(station_data, raw_data_text, report_date=None):
    """Use Claude to generate SAT-compliant XML from raw station data.

    Args:
        station_data: dict with station config (RFC, permit, tanks, etc.)
        raw_data_text: str with the raw operational data (readings, sales, etc.)
        report_date: date for the report (defaults to today)

    Returns:
        dict with xml_content, filename, filepath, and validation info
    """
    if not HAS_ANTHROPIC:
        return {"error": "Anthropic SDK not installed. Run: pip install anthropic"}

    if not ANTHROPIC_API_KEY:
        return {"error": "ANTHROPIC_API_KEY not configured. Set it in environment variables."}

    if report_date is None:
        report_date = date.today()

    # Build the user prompt with all context
    user_prompt = f"""Generate a complete SAT controles volumetricos DAILY XML report for this gas station.

STATION CONFIGURATION:
- RFC: {station_data.get('rfc', 'GAZ850101ABC')}
- RFC Proveedor Sistema: {station_data.get('rfc_proveedor', 'XAXX010101000')}
- Permiso: {station_data.get('num_permiso', 'PL/12345/EXP/ES/2024')}
- Modalidad: {station_data.get('modalidad_permiso', 'PL/XXXXX/EXP/ES/2024')}
- Clave Instalacion: {station_data.get('clave_instalacion', 'EDS-0001')}
- Descripcion: {station_data.get('descripcion', 'Estacion de servicio')}
- Latitud: {station_data.get('latitud', '31.6904')}
- Longitud: {station_data.get('longitud', '-106.4245')}
- Numero de Tanques: {station_data.get('num_tanques', '4')}
- Numero de Dispensarios: {station_data.get('num_dispensarios', '8')}

REPORT DATE: {report_date.isoformat()}

RAW OPERATIONAL DATA:
{raw_data_text}

Generate the COMPLETE XML with all products, tanks, dispensarios, recepciones, entregas, existencias, and bitacora. Make sure the volumetric balance validates for every tank."""

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

        message = client.messages.create(
            model="claude-opus-4-5-20251101",
            max_tokens=16000,
            system=SYSTEM_PROMPT,
            messages=[
                {"role": "user", "content": user_prompt}
            ]
        )

        xml_content = message.content[0].text.strip()

        # Clean up any markdown fences if present
        if xml_content.startswith("```"):
            lines = xml_content.split("\n")
            xml_content = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])

        # Validate XML is well-formed
        validation = validate_xml(xml_content)

        if not validation["valid"]:
            return {
                "error": f"Generated XML is not well-formed: {validation['error']}",
                "xml_content": xml_content,
                "validation": validation,
            }

        # Save and zip
        rfc = station_data.get('rfc', 'GAZ850101ABC')
        clave = station_data.get('clave_instalacion', 'EDS-0001')
        date_str = report_date.strftime('%Y%m%d')

        xml_filename = f"{rfc}_{clave}_{date_str}_DIA.xml"
        zip_filename = f"{rfc}_{clave}_{date_str}_DIA.xml.zip"

        report_dir = os.path.join(os.path.dirname(__file__), "generated_reports")
        os.makedirs(report_dir, exist_ok=True)

        xml_path = os.path.join(report_dir, xml_filename)
        zip_path = os.path.join(report_dir, zip_filename)

        # Write XML file
        with open(xml_path, "w", encoding="utf-8") as f:
            f.write(xml_content)

        # Create zip
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.write(xml_path, xml_filename)

        return {
            "success": True,
            "xml_filename": xml_filename,
            "zip_filename": zip_filename,
            "xml_path": xml_path,
            "zip_path": zip_path,
            "validation": validation,
            "tokens_used": {
                "input": message.usage.input_tokens,
                "output": message.usage.output_tokens,
            },
        }

    except anthropic.APIError as e:
        return {"error": f"Anthropic API error: {str(e)}"}
    except Exception as e:
        return {"error": f"Unexpected error: {str(e)}"}


def validate_xml(xml_content):
    """Validate that XML is well-formed and check volumetric balance."""
    result = {"valid": False, "error": None, "warnings": [], "products": []}

    try:
        root = ET.fromstring(xml_content)
        result["valid"] = True

        # Check namespaces
        covol_ns = "http://www.sat.gob.mx/ControlesVolumetricos"
        exp_ns = "http://www.sat.gob.mx/ControlesVolumetricos/Expendio"

        # Count products
        productos = root.findall(f"{{{covol_ns}}}PRODUCTO")
        result["product_count"] = len(productos)

        for prod in productos:
            clave = prod.get(f"{{{covol_ns}}}ClaveProducto", "unknown")
            marca = prod.get(f"{{{covol_ns}}}MarcaComercial", "")
            tanques = prod.findall(f"{{{covol_ns}}}TANQUE")
            dispensarios = prod.findall(f"{{{covol_ns}}}DISPENSARIO")

            prod_info = {
                "clave": clave,
                "marca": marca,
                "tanques": len(tanques),
                "dispensarios": len(dispensarios),
                "balance_ok": True,
            }

            # Check volumetric balance for each tank
            for tanque in tanques:
                existencias = tanque.find(f"{{{covol_ns}}}Existencias")
                if existencias is not None:
                    try:
                        anterior_el = existencias.find(f"{{{covol_ns}}}VolumenExistenciasAnterior")
                        recep_el = existencias.find(f"{{{covol_ns}}}VolumenAcumOpsRecep")
                        entreg_el = existencias.find(f"{{{covol_ns}}}VolumenAcumOpsEntreg")
                        final_el = existencias.find(f"{{{covol_ns}}}VolumenExistencias")

                        if all(el is not None for el in [anterior_el, recep_el, entreg_el, final_el]):
                            anterior = float(anterior_el.get(f"{{{covol_ns}}}ValorNumerico", "0"))
                            recep = float(recep_el.get(f"{{{covol_ns}}}ValorNumerico", "0"))
                            entreg = float(entreg_el.get(f"{{{covol_ns}}}ValorNumerico", "0"))
                            final = float(final_el.get(f"{{{covol_ns}}}ValorNumerico", "0"))

                            expected = anterior + recep - entreg
                            if abs(expected - final) > 1.0:  # 1L tolerance
                                clave_tq = tanque.get(f"{{{covol_ns}}}ClaveTanque", "?")
                                result["warnings"].append(
                                    f"Balance mismatch in {clave_tq}: {anterior}+{recep}-{entreg}={expected} vs {final}"
                                )
                                prod_info["balance_ok"] = False
                    except (ValueError, TypeError):
                        pass

            result["products"].append(prod_info)

        # Check bitacora
        bitacoras = root.findall(f"{{{covol_ns}}}BITACORA")
        result["bitacora_count"] = len(bitacoras)

        if not bitacoras:
            result["warnings"].append("BITACORA section is empty")

    except ET.ParseError as e:
        result["error"] = str(e)

    return result


def generate_demo_xml(report_date=None):
    """Generate a demo XML report using data from the database (no AI needed)."""
    from database import db, Station, FuelTransaction, InventorySnapshot

    if report_date is None:
        report_date = date.today()

    stations = Station.query.filter_by(active=True).order_by(Station.code).all()
    if not stations:
        return {"error": "No active stations found"}

    # Use first station as the primary
    station = stations[0]

    # Build raw data text from database
    raw_lines = [f"FECHA: {report_date.isoformat()}", f"ESTACION: {station.name} ({station.code})", ""]

    start = datetime.combine(report_date, datetime.min.time())
    end = datetime.combine(report_date, datetime.max.time())

    for ft in ["magna", "premium", "diesel"]:
        snap = InventorySnapshot.query.filter_by(
            station_id=station.id, fuel_type=ft, snapshot_date=report_date
        ).first()

        cap = getattr(station, f"{ft}_capacity", 40000)
        closing = snap.liters_on_hand if snap else 0

        received = float(db.session.query(
            db.func.coalesce(db.func.sum(FuelTransaction.liters), 0)
        ).filter(
            FuelTransaction.station_id == station.id,
            FuelTransaction.fuel_type == ft,
            FuelTransaction.transaction_type == "received",
            FuelTransaction.timestamp.between(start, end),
        ).scalar())

        sold = float(db.session.query(
            db.func.coalesce(db.func.sum(FuelTransaction.liters), 0)
        ).filter(
            FuelTransaction.station_id == station.id,
            FuelTransaction.fuel_type == ft,
            FuelTransaction.transaction_type == "sold",
            FuelTransaction.timestamp.between(start, end),
        ).scalar())

        opening = closing - received + sold

        raw_lines.append(f"TANQUE {ft.upper()}:")
        raw_lines.append(f"  Capacidad: {cap}L")
        raw_lines.append(f"  Inventario Inicial: {opening:.0f}L")
        raw_lines.append(f"  Litros Recibidos: {received:.0f}L")
        raw_lines.append(f"  Litros Vendidos: {sold:.0f}L")
        raw_lines.append(f"  Inventario Final: {closing:.0f}L")
        raw_lines.append("")

    raw_data_text = "\n".join(raw_lines)

    station_config = {
        "rfc": "GAZ850101ABC",
        "rfc_proveedor": "XAXX010101000",
        "num_permiso": "PL/12345/EXP/ES/2024",
        "modalidad_permiso": "PL/XXXXX/EXP/ES/2024",
        "clave_instalacion": "EDS-0001",
        "descripcion": f"{station.name}, {station.address}, {station.city}",
        "latitud": str(station.latitude or "31.6904"),
        "longitud": str(station.longitude or "-106.4245"),
        "num_tanques": "3",
        "num_dispensarios": "8",
    }

    return generate_sat_xml_with_ai(station_config, raw_data_text, report_date)
